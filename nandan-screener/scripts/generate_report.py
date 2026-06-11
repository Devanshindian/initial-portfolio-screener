#!/usr/bin/env python3
"""Minimal Markdown -> .docx converter for Nandan screening notes.

Renders in a strictly black-and-white house style (see the "Formatting Notes"
block in references/output-template.md): single neutral typeface in black,
plain Table Grid with thin black borders and no fill, sentence-case headings
sized by level, thin horizontal rules between sections, and a plain footer.

Supports: headings (#..####), paragraphs, bold (**..**), bullet lists (- / * ),
blockquotes (> ), horizontal rules (---), and GitHub-style pipe tables.
No pandoc dependency; uses python-docx only.
"""
import argparse, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)
FONT = "Calibri"
HEADING_SIZE = {1: 14, 2: 12, 3: 11, 4: 11}  # pt, by level


def add_runs(paragraph, text):
    # split on **bold**
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        else:
            paragraph.add_run(part)


def style_runs(paragraph, size=None, bold=None):
    # force every run to black Calibri, optionally size/bold
    for r in paragraph.runs:
        r.font.name = FONT
        r.font.color.rgb = BLACK
        if size is not None:
            r.font.size = Pt(size)
        if bold is not None:
            r.bold = bold


def add_hr(doc):
    # a real thin horizontal rule, not a row of underscores
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def extract_meta(md):
    # company name from the first H1, date from the first "Date:" line
    company, date = "", ""
    for line in md:
        s = line.strip()
        if not company and s.startswith("# "):
            company = s[2:].strip()
            company = re.sub(r"\s*-\s*Nandan Initial Screening Note$", "", company).strip()
        if not date and s.lower().startswith("date:"):
            date = s.split(":", 1)[1].strip()
        if company and date:
            break
    return company, date


def set_footer(doc, company, date):
    bits = ["Nandan Fund, Confidential"]
    if company:
        bits.append(company)
    bits.append("Initial Screening")
    if date:
        bits.append(date)
    fp = doc.sections[0].footer.paragraphs[0]
    fp.text = ". ".join(bits) + "."
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_runs(fp, size=8)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True)
    ap.add_argument("--output", required=True)
    a = ap.parse_args()
    md = open(a.input, encoding="utf-8").read().splitlines()

    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = FONT
    normal.size = Pt(11)
    normal.color.rgb = BLACK

    i = 0
    while i < len(md):
        line = md[i].rstrip()
        # table?
        if line.lstrip().startswith("|") and i + 1 < len(md) and is_table_sep(md[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows = []
            i += 2
            while i < len(md) and md[i].lstrip().startswith("|"):
                rows.append([c.strip() for c in md[i].strip().strip("|").split("|")])
                i += 1
            t = doc.add_table(rows=1, cols=len(header)); t.style = "Table Grid"
            for j, h in enumerate(header):
                cell = t.rows[0].cells[j]; cell.paragraphs[0].clear()
                add_runs(cell.paragraphs[0], h)
                style_runs(cell.paragraphs[0], size=10, bold=True)
            for row in rows:
                cells = t.add_row().cells
                for j in range(len(header)):
                    cells[j].paragraphs[0].clear()
                    add_runs(cells[j].paragraphs[0], row[j] if j < len(row) else "")
                    style_runs(cells[j].paragraphs[0], size=10)
            doc.add_paragraph()
            continue
        if not line.strip():
            i += 1; continue
        if line.startswith("#"):
            lvl = min(len(line) - len(line.lstrip("#")), 4)
            txt = line.lstrip("#").strip()
            h = doc.add_heading(level=lvl); add_runs(h, txt)
            style_runs(h, size=HEADING_SIZE[lvl], bold=True)
        elif line.strip() == "---":
            add_hr(doc)
        elif line.lstrip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet"); add_runs(p, line.lstrip()[2:])
            style_runs(p)
        elif line.lstrip().startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
            add_runs(p, line.lstrip()[1:].strip())
            style_runs(p); [setattr(r, "italic", True) for r in p.runs]
        else:
            p = doc.add_paragraph(); add_runs(p, line)
            style_runs(p)
        i += 1

    company, date = extract_meta(md)
    set_footer(doc, company, date)

    doc.save(a.output)
    print("Wrote", a.output)


if __name__ == "__main__":
    main()
